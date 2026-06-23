import csv
import io
import re
from datetime import datetime, timezone
from pathlib import Path
from uuid import uuid4

from bson import ObjectId
from fastapi import HTTPException, UploadFile, status
from openpyxl import load_workbook

from app.core.config import settings
from app.core.module_guard import ensure_tenant_module_enabled
from app.core.object_ids import parse_object_id, serialize_document
from app.core.permissions import get_owned_tenant_or_403
from app.db.mongodb import get_database
from app.services.rag_index_service import get_tenant_rag_status, reindex_tenant_knowledge
from app.services.rag_vector_service import delete_knowledge_document_vectors, upsert_knowledge_document_vectors

OWNER_SOURCE_TYPES = {"owner_text", "owner_upload"}
ALLOWED_UPLOAD_EXTENSIONS = {".txt", ".md", ".csv", ".xlsx", ".xlsm", ".pdf", ".docx"}
MAX_UPLOAD_SIZE_BYTES = 10 * 1024 * 1024
MAX_EXTRACTED_CHARS = 120000


def _clean_text(value: str) -> str:
    text = re.sub(r"\r\n?", "\n", str(value or ""))
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()[:MAX_EXTRACTED_CHARS]


def _normalize_tags(tags: list[str] | str | None) -> list[str]:
    if tags is None:
        return []
    raw_values = tags.split(",") if isinstance(tags, str) else tags
    normalized = []
    for tag in raw_values:
        value = str(tag or "").strip().lower()
        if value and value not in normalized:
            normalized.append(value[:40])
    return normalized[:20]


def _safe_file_name(file_name: str) -> str:
    name = Path(file_name or "knowledge-file").name
    name = re.sub(r"[^a-zA-Z0-9._ -]", "_", name).strip(" .")
    return name[:120] or "knowledge-file"


async def _ensure_knowledge_access(tenant_id: str, user: dict) -> tuple[ObjectId, dict]:
    tenant_oid = parse_object_id(tenant_id, "tenantId")
    tenant = await get_owned_tenant_or_403(tenant_oid, user)
    await ensure_tenant_module_enabled(tenant_oid, "ai_chat")
    return tenant_oid, tenant


def _serialize_document_summary(document: dict) -> dict:
    serialized = serialize_document(document)
    content = serialized.get("content") or ""
    serialized["contentPreview"] = content[:260]
    serialized["contentLength"] = len(content)
    serialized.pop("content", None)
    return serialized


async def _refresh_tenant_rag_summary(tenant_id: ObjectId) -> None:
    db = get_database()
    total = await db.knowledge_documents.count_documents({"tenantId": tenant_id})
    active = await db.knowledge_documents.count_documents({"tenantId": tenant_id, "isActive": True})
    chunks_cursor = db.knowledge_documents.find({"tenantId": tenant_id}, {"chunkCount": 1, "embeddingProvider": 1}).sort("lastIndexedAt", -1)
    chunk_count = 0
    provider = "none"
    async for row in chunks_cursor:
        chunk_count += int(row.get("chunkCount") or 0)
        if provider == "none" and row.get("embeddingProvider"):
            provider = row.get("embeddingProvider")
    await db.tenants.update_one(
        {"_id": tenant_id},
        {
            "$set": {
                "rag.lastIndexedAt": datetime.now(timezone.utc),
                "rag.knowledgeDocumentCount": total,
                "rag.activeKnowledgeDocumentCount": active,
                "rag.chunkCount": chunk_count,
                "rag.embeddingProvider": provider,
                "rag.indexStatus": "ready",
            }
        },
    )


async def _index_and_store_vector_state(document: dict) -> dict:
    db = get_database()
    now = datetime.now(timezone.utc)
    vector_result = await upsert_knowledge_document_vectors(document)
    await db.knowledge_documents.update_one(
        {"_id": document["_id"]},
        {
            "$set": {
                "embeddingProvider": vector_result["embeddingProvider"],
                "chunkCount": vector_result["chunkCount"],
                "lastIndexedAt": now,
                "updatedAt": now,
            }
        },
    )
    await _refresh_tenant_rag_summary(document["tenantId"])
    return await db.knowledge_documents.find_one({"_id": document["_id"]})


def _extract_text_from_plain_bytes(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _extract_text_from_csv(data: bytes) -> str:
    text = _extract_text_from_plain_bytes(data)
    rows = []
    reader = csv.reader(io.StringIO(text))
    for index, row in enumerate(reader, start=1):
        if index > 5000:
            rows.append("...remaining rows skipped for knowledge-base size limit...")
            break
        rows.append(" | ".join(str(cell).strip() for cell in row if str(cell).strip()))
    return "\n".join(row for row in rows if row)


def _extract_text_from_excel(data: bytes) -> str:
    workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    sections = []
    try:
        for sheet in workbook.worksheets:
            sections.append(f"Sheet: {sheet.title}")
            for index, row in enumerate(sheet.iter_rows(values_only=True), start=1):
                if index > 5000:
                    sections.append("...remaining rows skipped for knowledge-base size limit...")
                    break
                values = [str(value).strip() for value in row if value is not None and str(value).strip()]
                if values:
                    sections.append(" | ".join(values))
    finally:
        workbook.close()
    return "\n".join(sections)


def _extract_text_from_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="PDF parsing dependency is missing. Install pypdf from backend/requirements.txt and restart the API.",
        ) from exc

    reader = PdfReader(io.BytesIO(data))
    pages = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"Page {index}:\n{text.strip()}")
    return "\n\n".join(pages)


def _extract_text_from_docx(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="DOCX parsing dependency is missing. Install python-docx from backend/requirements.txt and restart the API.",
        ) from exc

    document = Document(io.BytesIO(data))
    lines = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                lines.append(" | ".join(values))
    return "\n".join(lines)


def _extract_upload_text(data: bytes, extension: str) -> str:
    if extension in {".txt", ".md"}:
        return _extract_text_from_plain_bytes(data)
    if extension == ".csv":
        return _extract_text_from_csv(data)
    if extension in {".xlsx", ".xlsm"}:
        return _extract_text_from_excel(data)
    if extension == ".pdf":
        return _extract_text_from_pdf(data)
    if extension == ".docx":
        return _extract_text_from_docx(data)
    raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="Unsupported knowledge-base file type.")


def _store_uploaded_file(tenant_id: ObjectId, document_id: ObjectId, safe_name: str, data: bytes) -> dict:
    base_dir = Path(settings.local_upload_dir).resolve()
    target_dir = base_dir / "knowledge-base" / str(tenant_id) / str(document_id)
    target_dir.mkdir(parents=True, exist_ok=True)
    stored_name = f"{uuid4().hex}-{safe_name}"
    destination = target_dir / stored_name
    destination.write_bytes(data)
    relative_path = f"knowledge-base/{tenant_id}/{document_id}/{stored_name}"
    return {
        "provider": "local",
        "fileId": relative_path,
        "url": f"/uploads/{relative_path}",
        "sizeBytes": len(data),
    }


async def list_knowledge_documents(
    tenant_id: str,
    user: dict,
    search: str = "",
    source_type: str | None = None,
    active_only: bool = False,
    page: int = 1,
    limit: int = 20,
) -> dict:
    db = get_database()
    tenant_oid, tenant = await _ensure_knowledge_access(tenant_id, user)
    page = max(page, 1)
    limit = min(max(limit, 1), 100)
    query: dict = {"tenantId": tenant_oid}
    if source_type:
        query["sourceType"] = source_type
    if active_only:
        query["isActive"] = True
    if search:
        query["$or"] = [
            {"title": {"$regex": search, "$options": "i"}},
            {"content": {"$regex": search, "$options": "i"}},
            {"tags": {"$regex": search, "$options": "i"}},
            {"metadata.fileName": {"$regex": search, "$options": "i"}},
        ]

    total = await db.knowledge_documents.count_documents(query)
    cursor = db.knowledge_documents.find(query).sort("updatedAt", -1).skip((page - 1) * limit).limit(limit)
    documents = [_serialize_document_summary(document) async for document in cursor]
    owner_upload_count = await db.knowledge_documents.count_documents({"tenantId": tenant_oid, "sourceType": "owner_upload"})
    owner_text_count = await db.knowledge_documents.count_documents({"tenantId": tenant_oid, "sourceType": "owner_text"})
    system_count = await db.knowledge_documents.count_documents({"tenantId": tenant_oid, "sourceType": {"$nin": list(OWNER_SOURCE_TYPES)}})

    return {
        "tenant": serialize_document(tenant),
        "items": documents,
        "summary": {
            "ownerUploadCount": owner_upload_count,
            "ownerTextCount": owner_text_count,
            "systemDocumentCount": system_count,
        },
        "pagination": {
            "page": page,
            "limit": limit,
            "total": total,
            "totalPages": (total + limit - 1) // limit,
        },
    }


async def get_knowledge_document(tenant_id: str, document_id: str, user: dict) -> dict:
    db = get_database()
    tenant_oid, tenant = await _ensure_knowledge_access(tenant_id, user)
    document_oid = parse_object_id(document_id, "documentId")
    document = await db.knowledge_documents.find_one({"_id": document_oid, "tenantId": tenant_oid})
    if not document:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Knowledge document not found.")
    return {"tenant": serialize_document(tenant), "document": serialize_document(document)}


async def create_text_knowledge_document(tenant_id: str, payload, user: dict) -> dict:
    db = get_database()
    tenant_oid, tenant = await _ensure_knowledge_access(tenant_id, user)
    now = datetime.now(timezone.utc)
    document_id = ObjectId()
    content = _clean_text(payload.content)
    if not content:
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="Knowledge content cannot be empty.")

    document = {
        "_id": document_id,
        "tenantId": tenant_oid,
        "branchId": None,
        "moduleCode": payload.moduleCode or "ai_chat",
        "sourceType": "owner_text",
        "sourceId": document_id,
        "title": payload.title.strip(),
        "content": content,
        "chromaCollection": f"tenant_{tenant_oid}",
        "chromaDocumentId": f"owner_text_{document_id}",
        "metadata": payload.metadata or {},
        "tags": _normalize_tags(payload.tags),
        "isActive": payload.isActive,
        "createdByUserId": user["_id"],
        "updatedByUserId": user["_id"],
        "embeddingProvider": "none",
        "chunkCount": 0,
        "lastIndexedAt": None,
        "createdAt": now,
        "updatedAt": now,
    }
    await db.knowledge_documents.insert_one(document)
    document = await _index_and_store_vector_state(document)
    return {"tenant": serialize_document(tenant), "document": serialize_document(document)}


async def upload_knowledge_document(
    tenant_id: str,
    file: UploadFile,
    user: dict,
    title: str | None = None,
    module_code: str = "ai_chat",
    tags: str | list[str] | None = None,
    is_active: bool = True,
) -> dict:
    db = get_database()
    tenant_oid, tenant = await _ensure_knowledge_access(tenant_id, user)
    safe_name = _safe_file_name(file.filename or "knowledge-file")
    extension = Path(safe_name).suffix.lower()
    if extension not in ALLOWED_UPLOAD_EXTENSIONS:
        allowed = ", ".join(sorted(ALLOWED_UPLOAD_EXTENSIONS))
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail=f"Unsupported file type. Allowed: {allowed}.")

    data = await file.read()
    if not data:
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="Uploaded file is empty.")
    if len(data) > MAX_UPLOAD_SIZE_BYTES:
        raise HTTPException(status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE, detail="Knowledge-base file must be 10MB or smaller.")

    extracted_text = _clean_text(_extract_upload_text(data, extension))
    if not extracted_text:
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="No readable text found in this file.")

    now = datetime.now(timezone.utc)
    document_id = ObjectId()
    stored = _store_uploaded_file(tenant_oid, document_id, safe_name, data)
    document_title = (title or Path(safe_name).stem or "Knowledge upload").strip()[:180]
    document = {
        "_id": document_id,
        "tenantId": tenant_oid,
        "branchId": None,
        "moduleCode": module_code or "ai_chat",
        "sourceType": "owner_upload",
        "sourceId": document_id,
        "title": document_title,
        "content": extracted_text,
        "chromaCollection": f"tenant_{tenant_oid}",
        "chromaDocumentId": f"owner_upload_{document_id}",
        "metadata": {
            "fileName": safe_name,
            "fileExtension": extension,
            "contentType": file.content_type or "",
            "storage": stored,
        },
        "tags": _normalize_tags(tags),
        "isActive": is_active,
        "createdByUserId": user["_id"],
        "updatedByUserId": user["_id"],
        "embeddingProvider": "none",
        "chunkCount": 0,
        "lastIndexedAt": None,
        "createdAt": now,
        "updatedAt": now,
    }
    await db.knowledge_documents.insert_one(document)
    document = await _index_and_store_vector_state(document)
    return {"tenant": serialize_document(tenant), "document": serialize_document(document)}


async def update_knowledge_document(tenant_id: str, document_id: str, payload, user: dict) -> dict:
    db = get_database()
    tenant_oid, tenant = await _ensure_knowledge_access(tenant_id, user)
    document_oid = parse_object_id(document_id, "documentId")
    document = await db.knowledge_documents.find_one({"_id": document_oid, "tenantId": tenant_oid})
    if not document:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Knowledge document not found.")
    if document.get("sourceType") not in OWNER_SOURCE_TYPES:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="System-generated knowledge cannot be edited here. Update the source module and reindex instead.")

    update = {"updatedAt": datetime.now(timezone.utc), "updatedByUserId": user["_id"]}
    should_reindex = False
    if payload.title is not None:
        update["title"] = payload.title.strip()
        should_reindex = True
    if payload.content is not None:
        content = _clean_text(payload.content)
        if not content:
            raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="Knowledge content cannot be empty.")
        update["content"] = content
        should_reindex = True
    if payload.moduleCode is not None:
        update["moduleCode"] = payload.moduleCode or "ai_chat"
        should_reindex = True
    if payload.tags is not None:
        update["tags"] = _normalize_tags(payload.tags)
    if payload.metadata is not None:
        existing_metadata = document.get("metadata") or {}
        update["metadata"] = {**existing_metadata, **(payload.metadata or {})}
    if payload.isActive is not None:
        update["isActive"] = payload.isActive
        should_reindex = True

    await db.knowledge_documents.update_one({"_id": document_oid, "tenantId": tenant_oid}, {"$set": update})
    updated = await db.knowledge_documents.find_one({"_id": document_oid, "tenantId": tenant_oid})
    if should_reindex:
        updated = await _index_and_store_vector_state(updated)
    else:
        await _refresh_tenant_rag_summary(tenant_oid)
    return {"tenant": serialize_document(tenant), "document": serialize_document(updated)}


async def delete_knowledge_document(tenant_id: str, document_id: str, user: dict) -> dict:
    db = get_database()
    tenant_oid, tenant = await _ensure_knowledge_access(tenant_id, user)
    document_oid = parse_object_id(document_id, "documentId")
    document = await db.knowledge_documents.find_one({"_id": document_oid, "tenantId": tenant_oid})
    if not document:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Knowledge document not found.")
    if document.get("sourceType") not in OWNER_SOURCE_TYPES:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="System-generated knowledge cannot be deleted here. Update the source module and reindex instead.")

    delete_knowledge_document_vectors(document)
    await db.knowledge_documents.delete_one({"_id": document_oid, "tenantId": tenant_oid})
    storage = ((document.get("metadata") or {}).get("storage") or {})
    file_id = storage.get("fileId")
    if file_id:
        try:
            (Path(settings.local_upload_dir).resolve() / file_id).unlink(missing_ok=True)
        except Exception:
            pass
    await _refresh_tenant_rag_summary(tenant_oid)
    return {"tenant": serialize_document(tenant), "deletedDocumentId": str(document_oid)}


async def reindex_knowledge_base_for_owner(tenant_id: str, user: dict) -> dict:
    tenant_oid, tenant = await _ensure_knowledge_access(tenant_id, user)
    result = await reindex_tenant_knowledge(tenant_oid)
    status_payload = await get_tenant_rag_status(tenant_id, user)
    status_payload["rag"].update(result)
    return {"tenant": serialize_document(tenant), "rag": status_payload["rag"]}
